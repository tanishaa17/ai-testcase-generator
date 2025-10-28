"""
Export Manager Module
Handles export of test cases in multiple formats for integration with external systems.
"""

import json
import os
from typing import List, Dict, Any, Optional
from datetime import datetime
import xml.etree.ElementTree as ET
import docx
from docx.shared import Pt

class ExportManager:
    """Manages export of test cases in various formats for enterprise integration."""
    
    def __init__(self):
        self.supported_formats = ["json", "gherkin", "xml", "excel", "docx", "pdf"]
    
    def export(self, test_cases: List[Dict[str, Any]], format: str, 
               output_path: Optional[str] = None) -> str:
        """
        Export test cases to specified format.
        
        Args:
            test_cases: List of test case dictionaries
            format: Export format (json, gherkin, xml, excel, docx, pdf)
            output_path: Optional output file path
        
        Returns:
            Exported content as string or file path
        """
        if format not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format}. Supported: {self.supported_formats}")
        
        if format == "json":
            return self._export_json(test_cases, output_path)
        elif format == "gherkin":
            return self._export_gherkin(test_cases, output_path)
        elif format == "xml":
            return self._export_xml(test_cases, output_path)
        elif format == "docx":
            return self._export_docx(test_cases, output_path)
        else:
            raise NotImplementedError(f"Format {format} not yet implemented")
    
    def _export_json(self, test_cases: List[Dict[str, Any]], output_path: Optional[str]) -> str:
        """Export as JSON."""
        content = json.dumps({
            "test_cases": test_cases,
            "export_timestamp": datetime.now().isoformat(),
            "total_tests": len(test_cases)
        }, indent=2)
        
        if output_path:
            with open(output_path, 'w') as f:
                f.write(content)
            return output_path
        return content
    
    def _export_gherkin(self, test_cases: List[Dict[str, Any]], output_path: Optional[str]) -> str:
        """Export as Gherkin feature files."""
        gherkin_content = []
        for tc in test_cases:
            gherkin_content.append(tc.get('gherkin_feature', ''))
        
        content = "\n\n".join(gherkin_content)
        
        if output_path:
            with open(output_path, 'w') as f:
                f.write(content)
            return output_path
        return content
    
    def _export_xml(self, test_cases: List[Dict[str, Any]], output_path: Optional[str]) -> str:
        """Export as XML for requirements tools."""
        root = ET.Element("testcases")
        root.set("version", "1.0")
        root.set("timestamp", datetime.now().isoformat())
        
        for tc in test_cases:
            tc_elem = ET.SubElement(root, "testcase")
            tc_elem.set("id", tc.get('test_id', ''))
            
            ET.SubElement(tc_elem, "requirement").text = tc.get('requirement_source', '')
            ET.SubElement(tc_elem, "gherkin").text = tc.get('gherkin_feature', '')
            
            # Compliance
            compliance_elem = ET.SubElement(tc_elem, "compliance")
            compliance_elem.set("status", tc.get('compliance_assessment', {}).get('status', 'Unknown'))
            
            tags_elem = ET.SubElement(compliance_elem, "tags")
            for tag in tc.get('compliance_tags', []):
                ET.SubElement(tags_elem, "tag").text = tag
            
            # Risk
            risk_elem = ET.SubElement(tc_elem, "risk")
            risk_elem.set("score", str(tc.get('risk_and_priority', {}).get('score', 0)))
        
        tree = ET.ElementTree(root)
        
        if output_path:
            tree.write(output_path, encoding='utf-8', xml_declaration=True)
            return output_path
        
        # Return as string
        ET.indent(tree, space="  ")
        return ET.tostring(root, encoding='unicode')
    
    def _export_docx(self, test_cases: List[Dict[str, Any]], output_path: Optional[str]) -> str:
        """Export as Word document with proper formatting."""
        doc = docx.Document()
        
        # Title
        title = doc.add_heading('Test Cases Export', 0)
        
        # Metadata
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Total Test Cases: {len(test_cases)}')
        
        for tc in test_cases:
            doc.add_page_break()
            
            # Test Case Header
            doc.add_heading(f"Test ID: {tc.get('test_id')}", level=1)
            
            # Requirement Source
            doc.add_paragraph('Requirement Source:', style='Heading 2')
            doc.add_paragraph(tc.get('requirement_source', ''))
            
            # Gherkin
            doc.add_paragraph('Gherkin Feature:', style='Heading 2')
            doc.add_paragraph(tc.get('gherkin_feature', ''), style='Preformatted')
            
            # Compliance
            doc.add_paragraph('Compliance Assessment:', style='Heading 2')
            compliance = tc.get('compliance_assessment', {})
            doc.add_paragraph(f"Status: {compliance.get('status', 'Unknown')}")
            doc.add_paragraph(f"Reasoning: {compliance.get('reasoning', '')}")
            
            # Tags
            tags = tc.get('compliance_tags', [])
            if tags:
                doc.add_paragraph('Compliance Tags:', style='Heading 2')
                doc.add_paragraph(', '.join(tags))
            
            # Risk
            risk = tc.get('risk_and_priority', {})
            doc.add_paragraph('Risk Assessment:', style='Heading 2')
            doc.add_paragraph(f"Score: {risk.get('score', 0)}/10")
            doc.add_paragraph(f"Reasoning: {risk.get('reasoning', '')}")
        
        if output_path:
            doc.save(output_path)
            return output_path
        
        # For string return, save to temp file
        temp_path = f"temp_export_{int(datetime.now().timestamp())}.docx"
        doc.save(temp_path)
        return temp_path

